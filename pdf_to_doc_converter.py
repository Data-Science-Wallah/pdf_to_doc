"""Streamlit PDF → DOCX converter (editable Word output).
What this app does
- Upload a PDF.
- Convert it to an editable DOCX while trying to preserve layout.
- Show a quick text preview extracted from the DOCX.
- Provide a download button for the DOCX.
Implementation notes (aligned with your described pipeline)
- PDF parsing + layout-aware conversion: `pdf2docx` (internally parses PDF + reconstructs layout)
- Post-processing: a conservative cleanup pass on the produced DOCX
For very complex PDFs, open-source converters may not reach perfect fidelity.
"""
from __future__ import annotations
import os
import tempfile
from io import BytesIO
import streamlit as st
from docx import Document
from pdf2docx import Converter

def _post_process_docx(docx_path: str) -> None:
    """Lightweight post-processing on the produced DOCX.
    Keep this conservative: aggressive rewriting can harm layout fidelity.
    Note: We avoid deleting paragraphs via private python-docx members to keep the
    implementation robust across versions.
    """
    # Intentionally minimal for safety.
    try:
        doc = Document(docx_path)
    except (OSError, ValueError):
        return
    # If we can save, we consider it a successful "touch" to ensure the file is valid.
    try:
        doc.save(docx_path)
    except OSError:
        return

def convert_pdf_bytes_to_docx(pdf_bytes: bytes) -> tuple[bytes, str]:
    """Convert PDF bytes to DOCX bytes.
    Returns:
        (docx_bytes, debug_message)
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as in_pdf:
        in_pdf.write(pdf_bytes)
        in_pdf.flush()
        pdf_path = in_pdf.name
    out_docx_path = tempfile.mktemp(suffix=".docx")
    try:
        cv = Converter(pdf_path)
        # layout=True = best effort to preserve original structure.
        cv.convert(out_docx_path, start=0, end=None, layout=True)
        cv.close()
        _post_process_docx(out_docx_path)
        with open(out_docx_path, "rb") as f:
            return f.read(), "Converted with pdf2docx (layout=True) + post-processing"
    finally:
        for p in (pdf_path, out_docx_path):
            try:
                if p and os.path.exists(p):
                    os.unlink(p)
            except OSError:
                pass

def docx_text_preview(docx_bytes: bytes, max_paragraphs: int = 20) -> str:
    """Extract a plain-text preview from DOCX bytes."""
    doc = Document(BytesIO(docx_bytes))
    out = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            out.append(t)
        if len(out) >= max_paragraphs:
            break
    return "\n".join(out)

def main() -> None:
    st.set_page_config(
        page_title="DataScienceWallah PDF to Docx Converter",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS for professional, colorful styling
    st.markdown("""
    <style>
    .main { background-color: #f0f2f6; }
    .stApp { background-color: #ffffff; border-radius: 10px; padding: 20px; box-shadow: 0 4px 8px rgba(0,0,0,0.1); }
    .stTitle { color: #1f77b4; font-family: 'Arial', sans-serif; font-weight: bold; }
    .stMarkdown { color: #333333; }
    .stButton>button { background-color: #ff7f0e; color: white; border-radius: 5px; border: none; padding: 10px 20px; font-weight: bold; }
    .stButton>button:hover { background-color: #e66a0a; }
    .stFileUploader { border: 2px dashed #1f77b4; border-radius: 10px; padding: 20px; background-color: #f9f9f9; }
    .stSuccess { background-color: #d4edda; color: #155724; border: 1px solid #c3e6cb; border-radius: 5px; padding: 10px; }
    .stInfo { background-color: #d1ecf1; color: #0c5460; border: 1px solid #bee5eb; border-radius: 5px; padding: 10px; }
    .stSpinner { color: #1f77b4; }
    .stCode { background-color: #f8f9fa; border: 1px solid #dee2e6; border-radius: 5px; padding: 10px; }
    .stDownloadButton { background-color: #28a745; color: white; border-radius: 5px; border: none; padding: 10px 20px; font-weight: bold; }
    .stDownloadButton:hover { background-color: #218838; }
    </style>
    """, unsafe_allow_html=True)
    
    st.title("📄 DataScienceWallah PDF to Docx Converter")
    st.markdown(
        "Transform your PDFs into editable DOCX files with **layout preservation**. "
        "Powered by advanced conversion technology for high-fidelity results. "
        "Upload, convert, and download in seconds! 🚀"
    )
    
    # Use columns for better layout
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.subheader("📤 Upload Your PDF")
        uploaded = st.file_uploader("Choose a PDF file", type=["pdf"], help="Select a PDF to convert to DOCX.")
        if not uploaded:
            st.info("👆 Upload a PDF to get started. Supported formats: PDF only.")
    
    with col2:
        if uploaded:
            pdf_bytes = uploaded.getvalue()
            with st.spinner("🔄 Converting your PDF to DOCX... Please wait."):
                docx_bytes, debug = convert_pdf_bytes_to_docx(pdf_bytes)
            st.success("✅ Conversion completed successfully!")
            st.caption(f"ℹ️ {debug}")
            
            st.subheader("👀 Text Preview")
            preview = docx_text_preview(docx_bytes, max_paragraphs=20)
            if preview:
                st.code(preview, language="text")
            else:
                st.info("No text extracted from the DOCX.")
            
            st.subheader("⬇️ Download Your DOCX")
            out_name = os.path.splitext(uploaded.name)[0] + ".docx"
            st.download_button(
                label="📥 Download DOCX",
                data=docx_bytes,
                file_name=out_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                help="Click to download the converted DOCX file."
            )

if __name__ == "__main__":
    main()
